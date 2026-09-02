#!/usr/bin/env python3
"""Build the manual E2E test plan as a PDF and a DOCX from its Markdown sources.

The sources live in docs/qa/src/*.md and are concatenated in filename order.
The Markdown subset used there is deliberately small — headings, tables, lists,
fenced code, blockquotes, horizontal rules, and inline bold/italic/code — so a
self-contained converter is cheaper than a dependency, and produces exactly the
print layout we want in both formats.

PDF: rendered to HTML with print CSS, then printed by the Chromium that ships
with Playwright (no Node modules are installed in this repo).
DOCX: built with python-docx from the same parsed blocks, so the two documents
can never describe different content.
"""

import html
import os
import re
import subprocess
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "qa"

# The two documents this script builds. Same parser, same CSS, same DOCX
# writer -- the only thing that differs is which source directory feeds it
# and what the cover says, so the plan and the audit of the plan can never
# drift into two different house styles.
DOCS = {
    "plan": {
        "src": OUT / "src",
        "basename": "DrPoojaPhysio-E2E-Test-Plan",
        "title": "Dr. Pooja's Physio — Complete Manual E2E Test Plan & Feature Guide",
        "subtitle": "Feature guide and click-by-click regression suite",
        "meta": [
            ("Application", "Dr. Pooja's Physio — Next.js 16 · React 19 · Supabase · Razorpay"),
            ("Document version", "1.0"),
            ("Audience", "A tester who has never used this application before"),
            ("Environment", "Throwaway Supabase project · Razorpay test mode · npm run dev"),
            ("Start here", "Section 6 — STEP 0, Reset the test environment"),
        ],
    },
    "audit": {
        "src": OUT / "audit-src",
        "basename": "DrPoojaPhysio-QA-Audit-Report",
        "title": "Dr. Pooja's Physio — QA Audit Report",
        "subtitle": "Static verification against the manual E2E test plan, plus a product review",
        "meta": [
            ("Subject", "Branch claude/complete-e2e-testing-plan-910y5z"),
            ("Method", "Executed checks plus source verification — the plan was NOT executed"),
            ("Verdict", "Conditional pass — 6 findings, 1 to fix before the first test run"),
            ("Read first", "Section 1 — Scope, and what this report is not"),
        ],
    },
}

CHROME_CANDIDATES = [
    "/opt/pw-browsers/chromium-1194/chrome-linux/chrome",
    "/opt/pw-browsers/chromium/chrome-linux/chrome",
    "/usr/bin/chromium",
    "/usr/bin/google-chrome",
]


# --------------------------------------------------------------------------
# Parsing: Markdown -> a flat list of blocks
# --------------------------------------------------------------------------

def read_source(src: Path) -> str:
    parts = sorted(p for p in src.glob("*.md"))
    if not parts:
        sys.exit(f"no markdown sources found in {src}")
    return "\n\n".join(p.read_text(encoding="utf-8").rstrip() for p in parts) + "\n"


def parse(md: str):
    """Returns a list of (kind, payload) blocks."""
    blocks = []
    lines = md.split("\n")
    i = 0
    para: list[str] = []

    def flush_para():
        nonlocal para
        if para:
            blocks.append(("p", " ".join(para).strip()))
            para = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # fenced code
        if stripped.startswith("```"):
            flush_para()
            i += 1
            code = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1
            blocks.append(("code", "\n".join(code)))
            continue

        if not stripped:
            flush_para()
            i += 1
            continue

        # horizontal rule
        if re.fullmatch(r"-{3,}", stripped):
            flush_para()
            blocks.append(("hr", ""))
            i += 1
            continue

        # heading
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            flush_para()
            blocks.append((f"h{len(m.group(1))}", m.group(2).strip()))
            i += 1
            continue

        # table: a header row followed by a separator row
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(
            r"^\|[\s:\-|]+\|$", lines[i + 1].strip()
        ):
            flush_para()
            rows = []
            header = split_row(stripped)
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i].strip()))
                i += 1
            blocks.append(("table", (header, rows)))
            continue

        # blockquote
        if stripped.startswith(">"):
            flush_para()
            quote = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote.append(lines[i].strip().lstrip(">").strip())
                i += 1
            blocks.append(("quote", " ".join(x for x in quote if x)))
            continue

        # lists
        if re.match(r"^\s*([-*]|\d+\.)\s+", line):
            flush_para()
            items = []
            ordered = bool(re.match(r"^\s*\d+\.\s+", line))
            while i < len(lines) and re.match(r"^\s*([-*]|\d+\.)\s+", lines[i]):
                indent = len(lines[i]) - len(lines[i].lstrip())
                text = re.sub(r"^\s*([-*]|\d+\.)\s+", "", lines[i]).strip()
                i += 1
                # fold continuation lines into the item
                while (
                    i < len(lines)
                    and lines[i].strip()
                    and not re.match(r"^\s*([-*]|\d+\.)\s+", lines[i])
                    and not lines[i].strip().startswith(("#", "|", ">", "```"))
                    and (len(lines[i]) - len(lines[i].lstrip())) > indent
                ):
                    text += " " + lines[i].strip()
                    i += 1
                items.append((indent // 2, text))
            blocks.append(("ol" if ordered else "ul", items))
            continue

        para.append(stripped)
        i += 1

    flush_para()
    return blocks


def split_row(row: str) -> list[str]:
    cells = row.strip().strip("|").split("|")
    return [c.strip() for c in cells]


# --------------------------------------------------------------------------
# Inline formatting
# --------------------------------------------------------------------------

INLINE_RE = re.compile(
    r"(`[^`]+`)"           # code
    r"|(\*\*[^*]+\*\*)"    # bold
    r"|(\*[^*]+\*)"        # italic
    r"|(\[[^\]]+\]\([^)]+\))"  # link
)


def inline_tokens(text: str):
    """Yields (style, text) where style is one of: plain, code, bold, italic, link."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            yield ("plain", text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("`"):
            yield ("code", tok[1:-1])
        elif tok.startswith("**"):
            yield ("bold", tok[2:-2])
        elif tok.startswith("["):
            label, href = re.match(r"\[([^\]]+)\]\(([^)]+)\)", tok).groups()
            yield ("link", (label, href))
        else:
            yield ("italic", tok[1:-1])
        pos = m.end()
    if pos < len(text):
        yield ("plain", text[pos:])


def inline_html(text: str) -> str:
    out = []
    for style, value in inline_tokens(text):
        if style == "link":
            label, href = value
            out.append(f'<a href="{html.escape(href)}">{html.escape(label)}</a>')
        elif style == "code":
            out.append(f"<code>{html.escape(value)}</code>")
        elif style == "bold":
            out.append(f"<strong>{html.escape(value)}</strong>")
        elif style == "italic":
            out.append(f"<em>{html.escape(value)}</em>")
        else:
            out.append(html.escape(value))
    return "".join(out)


# --------------------------------------------------------------------------
# HTML / PDF
# --------------------------------------------------------------------------

CSS = """
@page { size: A4; margin: 18mm 15mm 20mm 15mm; }
* { box-sizing: border-box; }
body {
  font-family: "DejaVu Sans", "Segoe UI", Helvetica, Arial, sans-serif;
  font-size: 9.4pt; line-height: 1.5; color: #1e293b; margin: 0;
}
.cover { page-break-after: always; padding-top: 55mm; text-align: center; }
.cover .rule { width: 60mm; height: 4px; background: #0f766e; margin: 0 auto 14mm; }
.cover h1 { font-size: 26pt; line-height: 1.2; color: #0f172a; margin: 0 0 10mm; }
.cover .sub { font-size: 12pt; color: #475569; margin-bottom: 22mm; }
.cover dl { display: inline-block; text-align: left; font-size: 10pt; color: #334155; }
.cover dt { font-weight: 700; color: #0f766e; margin-top: 5mm; }
.toc { page-break-after: always; }
.toc h2 { border: none; }
.toc ul { columns: 2; column-gap: 12mm; font-size: 9pt; padding-left: 0; list-style: none; }
.toc li { margin-bottom: 1.6mm; break-inside: avoid; }
h1 { font-size: 17pt; color: #0f172a; margin: 10mm 0 3mm; page-break-after: avoid; }
h2 {
  font-size: 13.5pt; color: #0f172a; margin: 9mm 0 3mm;
  padding-bottom: 1.5mm; border-bottom: 2px solid #0f766e; page-break-after: avoid;
}
h3 { font-size: 11.5pt; color: #0f766e; margin: 6mm 0 2mm; page-break-after: avoid; }
h4 {
  font-size: 10pt; color: #0f172a; margin: 5mm 0 1.5mm;
  page-break-after: avoid; page-break-inside: avoid;
}
h5, h6 { font-size: 9.6pt; color: #334155; margin: 4mm 0 1.5mm; }
p { margin: 0 0 2.4mm; }
ul, ol { margin: 0 0 2.6mm; padding-left: 6mm; }
li { margin-bottom: 1.1mm; }
code {
  font-family: "DejaVu Sans Mono", Consolas, monospace; font-size: 8.4pt;
  background: #f1f5f9; border: 1px solid #e2e8f0; border-radius: 3px; padding: 0 1.2mm;
  word-break: break-word;
}
pre {
  background: #0f172a; color: #e2e8f0; padding: 3mm 4mm; border-radius: 4px;
  font-family: "DejaVu Sans Mono", Consolas, monospace; font-size: 8pt;
  line-height: 1.45; overflow-x: auto; white-space: pre-wrap;
  page-break-inside: avoid; margin: 0 0 3mm;
}
pre code { background: none; border: none; color: inherit; padding: 0; font-size: inherit; }
blockquote {
  margin: 0 0 3mm; padding: 2.5mm 4mm; background: #fffbeb;
  border-left: 3px solid #d97706; color: #78350f; font-size: 9pt;
}
table {
  width: 100%; border-collapse: collapse; margin: 0 0 3.5mm;
  font-size: 8.3pt; page-break-inside: auto;
}
th, td {
  border: 1px solid #cbd5e1; padding: 1.4mm 1.8mm; text-align: left;
  vertical-align: top; word-break: normal; overflow-wrap: anywhere;
}
th { background: #0f766e; color: #fff; font-weight: 700; }
tr { page-break-inside: avoid; }
tbody tr:nth-child(even) { background: #f8fafc; }
hr { border: none; border-top: 1px solid #e2e8f0; margin: 6mm 0; }
a { color: #0f766e; text-decoration: none; word-break: break-all; }
strong { color: #0f172a; }
"""


def to_html(blocks, doc) -> str:
    body = []
    for kind, payload in blocks:
        if kind.startswith("h") and len(kind) == 2 and kind[1].isdigit():
            level = int(kind[1])
            body.append(f"<h{level}>{inline_html(payload)}</h{level}>")
        elif kind == "p":
            body.append(f"<p>{inline_html(payload)}</p>")
        elif kind == "hr":
            body.append("<hr>")
        elif kind == "quote":
            body.append(f"<blockquote>{inline_html(payload)}</blockquote>")
        elif kind == "code":
            body.append(f"<pre><code>{html.escape(payload)}</code></pre>")
        elif kind in ("ul", "ol"):
            tag = "ul" if kind == "ul" else "ol"
            items = "".join(f"<li>{inline_html(t)}</li>" for _, t in payload)
            body.append(f"<{tag}>{items}</{tag}>")
        elif kind == "table":
            header, rows = payload
            head = "".join(f"<th>{inline_html(c)}</th>" for c in header)
            trs = []
            for r in rows:
                cells = r + [""] * (len(header) - len(r))
                trs.append(
                    "<tr>" + "".join(f"<td>{inline_html(c)}</td>" for c in cells[: len(header)]) + "</tr>"
                )
            body.append(
                f"<table><thead><tr>{head}</tr></thead><tbody>{''.join(trs)}</tbody></table>"
            )

    # Cover + table of contents from the level-2 headings.
    sections = [p for k, p in blocks if k == "h2"]
    toc_items = "".join(f"<li>{inline_html(s)}</li>" for s in sections)

    meta = "".join(
        f"<dt>{html.escape(k)}</dt><dd>{html.escape(v)}</dd>" for k, v in doc["meta"]
    )
    cover = f"""
<div class="cover">
  <div class="rule"></div>
  <h1>{html.escape(doc["title"])}</h1>
  <p class="sub">{html.escape(doc["subtitle"])}</p>
  <dl>{meta}</dl>
</div>
<div class="toc">
  <h2>Contents</h2>
  <ul>{toc_items}</ul>
</div>
"""

    return (
        "<!doctype html><html><head><meta charset='utf-8'>"
        f"<title>{html.escape(doc['title'])}</title><style>{CSS}</style></head><body>"
        + cover
        + "\n".join(body)
        + "</body></html>"
    )


def write_pdf(html_path: Path, pdf_path: Path) -> None:
    chrome = next((c for c in CHROME_CANDIDATES if os.path.exists(c)), None)
    if not chrome:
        print("! no Chromium found; skipping PDF", file=sys.stderr)
        return
    subprocess.run(
        [
            chrome,
            "--headless",
            "--disable-gpu",
            "--no-sandbox",
            "--no-pdf-header-footer",
            f"--print-to-pdf={pdf_path}",
            str(html_path),
        ],
        check=True,
        capture_output=True,
    )


# --------------------------------------------------------------------------
# DOCX
# --------------------------------------------------------------------------

def write_docx(blocks, path: Path, doc) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    docx = Document()
    normal = docx.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    def add_runs(paragraph, text):
        for style, value in inline_tokens(text):
            if style == "link":
                label, href = value
                run = paragraph.add_run(f"{label} ({href})")
                run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
            else:
                run = paragraph.add_run(value)
                if style == "bold":
                    run.bold = True
                elif style == "italic":
                    run.italic = True
                elif style == "code":
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)

    # Cover
    title = docx.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(doc["title"])
    run.bold = True
    run.font.size = Pt(22)
    sub = docx.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(doc["subtitle"]).italic = True
    for label, value in doc["meta"]:
        p = docx.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)
    docx.add_page_break()

    for kind, payload in blocks:
        if kind in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = min(int(kind[1]), 4)
            heading = docx.add_heading(level=level)
            add_runs(heading, payload)
        elif kind == "p":
            add_runs(docx.add_paragraph(), payload)
        elif kind == "quote":
            p = docx.add_paragraph(style="Intense Quote")
            add_runs(p, payload)
        elif kind == "hr":
            docx.add_paragraph("_" * 60)
        elif kind == "code":
            p = docx.add_paragraph()
            run = p.add_run(payload)
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)
        elif kind in ("ul", "ol"):
            style = "List Bullet" if kind == "ul" else "List Number"
            for depth, text in payload:
                name = style if depth == 0 else f"{style} {min(depth + 1, 3)}"
                try:
                    p = docx.add_paragraph(style=name)
                except KeyError:
                    p = docx.add_paragraph(style=style)
                add_runs(p, text)
        elif kind == "table":
            header, rows = payload
            table = docx.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            for cell, text in zip(table.rows[0].cells, header):
                cell.text = ""
                add_runs(cell.paragraphs[0], text)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for r in rows:
                cells = (r + [""] * len(header))[: len(header)]
                row = table.add_row()
                for cell, text in zip(row.cells, cells):
                    cell.text = ""
                    add_runs(cell.paragraphs[0], text)
            docx.add_paragraph()

    docx.save(path)


def build(doc) -> None:
    md = read_source(doc["src"])
    blocks = parse(md)

    OUT.mkdir(parents=True, exist_ok=True)
    base = doc["basename"]
    md_path = OUT / f"{base}.md"
    html_path = OUT / f"{base}.html"
    pdf_path = OUT / f"{base}.pdf"
    docx_path = OUT / f"{base}.docx"

    md_path.write_text(md, encoding="utf-8")
    html_path.write_text(to_html(blocks, doc), encoding="utf-8")
    write_pdf(html_path, pdf_path)
    write_docx(blocks, docx_path, doc)

    for p in (md_path, html_path, pdf_path, docx_path):
        if p.exists():
            print(f"{p.relative_to(ROOT)}  {p.stat().st_size / 1024:.0f} KB")


def main() -> None:
    # No argument builds both; a name builds one.
    wanted = sys.argv[1:] or list(DOCS)
    for name in wanted:
        if name not in DOCS:
            sys.exit(f"unknown document {name!r}; choose from {', '.join(DOCS)}")
        build(DOCS[name])


if __name__ == "__main__":
    main()
